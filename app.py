# Run this Streamlit app with: streamlit run app.py

import streamlit as st
import docx
import tempfile
import openai
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document

# === Configuration ===
openai.api_key = st.secrets["OPENAI_API_KEY"]

# === Helper Functions ===
def extract_text_from_docx(doc_file):
    doc = docx.Document(doc_file)
    return "\n".join([para.text for para in doc.paragraphs])

def compute_ats_score(resume_text, job_desc):
    vectorizer = CountVectorizer().fit_transform([resume_text, job_desc])
    vectors = vectorizer.toarray()
    score = cosine_similarity([vectors[0]], [vectors[1]])[0][0]
    return round(score * 100, 2)

def tailor_with_llm(resume_text, job_desc):
    prompt = f"""
You are a professional resume editor. Modify the following resume to better match the given job description by:
- Highlighting relevant skills and experiences
- Rewriting descriptions to emphasize job-specific keywords
- Keeping formatting and tone professional

Resume:
{resume_text}

Job Description:
{job_desc}

Tailored Resume:
"""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert resume optimization assistant."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.7
    )
    return response["choices"][0]["message"]["content"].strip()

def save_as_docx(text):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    temp_path = tempfile.mktemp(suffix=".docx")
    doc.save(temp_path)
    return temp_path

# === Streamlit App ===
st.set_page_config(page_title="Resume Tailor AI", layout="wide")
st.title("🧠 Resume Tailoring Assistant")

uploaded_file = st.file_uploader("Upload your Resume (DOCX format)", type=["docx"])
job_description = st.text_area("Paste the Job Description", height=200)

if uploaded_file and job_description:
    with st.spinner("Reading resume..."):
        resume_text = extract_text_from_docx(uploaded_file)

    pre_score = compute_ats_score(resume_text, job_description)
    st.markdown(f"### ATS Score Before Tailoring: {pre_score} ✅")

    if st.button("Tailor My Resume with AI"):
        with st.spinner("Tailoring resume using LLM..."):
            tailored_resume = tailor_with_llm(resume_text, job_description)
            post_score = compute_ats_score(tailored_resume, job_description)
            docx_path = save_as_docx(tailored_resume)

        st.markdown(f"### ATS Score After Tailoring: {post_score} 🚀")
        st.download_button("Download Tailored Resume (DOCX)", open(docx_path, "rb"), file_name="tailored_resume.docx")

        st.markdown("### Tailored Resume Preview")
        st.text_area("", value=tailored_resume, height=500)

st.markdown("\n---\nTo run this app locally: `streamlit run app.py`")
